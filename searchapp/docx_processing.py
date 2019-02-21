#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Functions for extracting information from Word documents
"""

from __future__ import print_function, division

import re
from datetime import datetime
from pprint import pprint  # useful for troubleshooting dictionary content

from docx import Document


def element_extract(element):
    """Element extract will be our ultimate utility function for both
    paragraph_extract and header_extract. An "element" in this case could be a
    document, or the content of the cells in a table.

    You need to return the raw text, where all newlines are replaced with
    spaces.

    Args:
        element: any component of a DOCX that has the paragraphs property.

    Returns:
        string of text with only single spaces splitting words (all newlines
            and multiple spaces are removed)
    """
    
    text = []
    for p in element.paragraphs:
        #print(p.text)
        newwords = p.text.split()
        if len(newwords) > 0:
            text.append(' '.join(newwords))
            
    return " ".join(text)


def paragraph_extract(inputFile):
    """This uses the element_extract function and passes the entire document.
    It does not require any modification.

    Args:
        inputFile (file): the DOCX file to be read

    Returns:
        string of text from the document skipping word art, tables, etc.
    """

    doc = Document(inputFile)
    return element_extract(doc)


def header_extract(inputFile, splitFields=[]):
    """From our schema we can assume that the first table is the only table we
    need to care about in the document as the header and all other tables can be
    skipped for now.

    The function takes an optional argument of splitFields which is an array of
    potential values. For each header key that is listed in splitFields the data
    should be split into an array rather than read as a single string. The array
    can be comma or space delimited so both should be accounted for.

    The header should return a dictionary, d, with key-value pairs for each
    entry in the table and the addition of a new key for date formatted as YYYY-
    MM-DD. This date will act as our timestamp for when the document is loaded
    in the search index. It will be set to "today".

    Args:
        inputFile (file): the DOCX file to be read
        splitFields (list): array of header keys that should be treated as
            array elements and not as a single string.

    Returns:
        dictionary of header-key: value for each entry in the first table of the
            word document or a blank dictionary if none exists. The dictionary
            should add a key for the date formatted YYYY-MM-DD.
    """
    
    headers = {}
    doc = Document(inputFile)
    if len(doc.tables) > 0:
        for r in doc.tables[0].rows:
            key = r.cells[0].text.lower()
            value = r.cells[1].text
            if key in splitFields:
                value = re.compile('\w+').findall(value)
            headers[key] = value
            #print(key, value)
        headers['date'] = datetime.today().strftime('%Y-%m-%d')
    return headers


if __name__ == '__main__':
    text = paragraph_extract('transcriptions/transcript01.docx')
    print(text)
    headers = header_extract('transcriptions/transcript01.docx', ['Tags'])
    print(headers)
